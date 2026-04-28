"""Generic markdown → DOCX renderer used by the four corrected papers.

Handles: # H1 (centered title), ## H1, ### H2, #### H3 (treated as H2),
- bullets, 1. numbered, **bold** *italic* `code` inline,
markdown tables, --- horizontal rules, italic-only paragraphs (subtitle / captions).
Optionally appends figures listed with --figure args.
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TITLE_COLOR = RGBColor(0x1A, 0x47, 0x73)
ACCENT = RGBColor(0x55, 0x55, 0x55)


def make_doc() -> Document:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.25
    for lv in range(1, 4):
        hs = doc.styles[f"Heading {lv}"]
        hs.font.color.rgb = TITLE_COLOR; hs.font.name = "Calibri"
    return doc


def add_rich(doc, text):
    p = doc.add_paragraph()
    for part in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", text):
        if not part: continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Menlo"; r.font.size = Pt(10)
        else:
            p.add_run(part)


def add_inline_runs(p, text):
    for part in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", text):
        if not part: continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Menlo"; r.font.size = Pt(10)
        else:
            p.add_run(part)


def _split_md_row(line: str) -> list[str]:
    """Split a markdown table row on `|`, respecting `\\|` escapes."""
    s = line.strip().strip("|")
    cells = []
    cur = []
    i = 0
    while i < len(s):
        if s[i] == "\\" and i + 1 < len(s) and s[i + 1] == "|":
            cur.append("|")
            i += 2
        elif s[i] == "|":
            cells.append("".join(cur).strip())
            cur = []
            i += 1
        else:
            cur.append(s[i])
            i += 1
    cells.append("".join(cur).strip())
    return cells


def add_table(doc, rows):
    if not rows or len(rows) < 2: return
    parsed = []
    for line in rows:
        if re.match(r"^\|\s*-", line): continue
        parsed.append(_split_md_row(line))
    if not parsed: return
    headers, body = parsed[0], parsed[1:]
    t = doc.add_table(rows=1+len(body), cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        for pr in t.rows[0].cells[j].paragraphs:
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in pr.runs: r.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(body):
        for j, cell in enumerate(row):
            if j >= len(headers): continue
            txt = re.sub(r"\*\*(.+?)\*\*", r"\1", cell)
            t.rows[i+1].cells[j].text = txt
            for pr in t.rows[i+1].cells[j].paragraphs:
                pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in pr.runs: r.font.size = Pt(9)
    doc.add_paragraph()


def render(md_path: Path, doc: Document, figure_queue: list[tuple[Path, str]] | None = None) -> None:
    """Render `md_path` into `doc`. If `figure_queue` is provided, italic-only
    paragraphs that begin with the word 'Figure' get the next queued image
    embedded immediately above them (centered, 6.5 in wide), and the caption
    text is taken from the figure-queue entry rather than the markdown line."""
    src = md_path.read_text()
    lines = src.split("\n")
    fig_queue = list(figure_queue) if figure_queue else []
    i = 0
    tbuf = []
    in_code = False
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            in_code = not in_code
            i += 1; continue
        if in_code:
            i += 1; continue
        if line.startswith("|") and "|" in line[1:]:
            tbuf.append(line); i += 1; continue
        elif tbuf:
            add_table(doc, tbuf); tbuf = []
        if line == "---":
            doc.add_paragraph("───────────────").alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("# "):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip()); r.bold = True
            r.font.size = Pt(20); r.font.color.rgb = TITLE_COLOR
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip())
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s", "", line))
        elif line == "":
            pass
        elif line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            caption_text = line.strip("*")
            # If this looks like a figure caption AND we have figures queued,
            # embed the next figure inline ABOVE the caption.
            if fig_queue and caption_text.lower().startswith("figure"):
                fig_path, fig_caption = fig_queue.pop(0)
                if fig_path.exists():
                    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pp.add_run().add_picture(str(fig_path), width=Inches(6.5))
                # Use the markdown caption (it's the in-context one); ignore
                # the --figure caption arg, which is just for ordering.
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(caption_text); r.italic = True; r.font.color.rgb = ACCENT
        else:
            add_rich(doc, line)
        i += 1
    if tbuf: add_table(doc, tbuf)
    # Any unused figures (no matching caption) get appended at the end.
    for fig_path, fig_caption in fig_queue:
        if not fig_path.exists():
            continue
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.add_run().add_picture(str(fig_path), width=Inches(6.5))
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(fig_caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = ACCENT


def add_figures(doc: Document, figures: list[tuple[Path, str]]) -> None:
    for fig_path, caption in figures:
        if not fig_path.exists():
            continue
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fig_path), width=Inches(6.5))
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = ACCENT


def main(argv=None) -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--input", required=True, help="Input markdown file")
    p.add_argument("--output", required=True, help="Output DOCX path")
    p.add_argument("--figure", action="append", default=[],
                   help="path:caption (can be repeated)")
    args = p.parse_args(argv)

    figures = []
    for f in args.figure:
        if ":" in f:
            path, _, caption = f.partition(":")
            figures.append((Path(path), caption))
        else:
            figures.append((Path(f), ""))
    doc = make_doc()
    render(Path(args.input), doc, figure_queue=figures)
    doc.save(args.output)
    size = Path(args.output).stat().st_size // 1024
    print(f"Wrote {args.output} ({size} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
